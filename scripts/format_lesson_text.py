from docx import Document
from docx.shared import RGBColor
import shutil
import os
from datetime import datetime

def format_specific_phrases(doc_path):
    """
    Word文書内の特定のフレーズを赤字太字に変更する
    バックアップを作成してから編集する
    """
    # バックアップを作成
    backup_path = doc_path.replace('.docx', f'_backup_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx')
    try:
        shutil.copy2(doc_path, backup_path)
        print(f"📋 バックアップ作成: {os.path.basename(backup_path)}")
    except Exception as e:
        print(f"⚠️ バックアップ作成失敗: {e}")
        print("元のファイルが開かれている可能性があります。閉じてから再実行してください。")
        return

    try:
        doc = Document(doc_path)

        # 検索する文字列
        target_phrases = [
            "目に酸素と潤いを",
            "口角を上げて笑顔で"
        ]

        found_count = 0

        # 全ての段落を処理
        for paragraph in doc.paragraphs:
            # 段落全体のテキストを取得
            para_text = paragraph.text

            # いずれかのフレーズが含まれているかチェック
            has_target = any(phrase in para_text for phrase in target_phrases)

            if has_target:
                # 段落のrunを再構築
                # 既存のrunをクリア
                for run in paragraph.runs:
                    run.text = ''

                # 段落を再構築
                remaining_text = para_text
                for phrase in target_phrases:
                    parts = remaining_text.split(phrase)
                    if len(parts) > 1:
                        found_count += len(parts) - 1
                        # 分割されたテキストを再結合
                        new_text = []
                        for i, part in enumerate(parts):
                            if i > 0:
                                new_text.append(('RED_BOLD', phrase))
                            new_text.append(('NORMAL', part))

                        # runを追加
                        for style, text in new_text:
                            if text:
                                run = paragraph.add_run(text)
                                if style == 'RED_BOLD':
                                    run.bold = True
                                    run.font.color.rgb = RGBColor(255, 0, 0)

                        remaining_text = ''.join([t for s, t in new_text if s != 'RED_BOLD'])
                        break

        # 保存
        output_path = doc_path.replace('.docx', '_formatted.docx')
        doc.save(output_path)
        print(f"\n✅ 完了: {os.path.basename(output_path)}")
        print(f"   - '目に酸素と潤いを' → 赤字太字")
        print(f"   - '口角を上げて笑顔で' → 赤字太字")
        print(f"   - {found_count}箇所を修正しました")
        print(f"\n📝 元のファイルを閉じて、このファイル名を変更してください")

    except Exception as e:
        print(f"❌ エラー: {e}")
        print("元のファイルが開かれている可能性があります。閉じてから再実行してください。")

if __name__ == "__main__":
    doc_path = r"C:\Users\info\Desktop\yoshiyuki\古事記project\VIBE CODING\sokudoku\資料\レッスン台本.docx"

    print("=" * 60)
    print("レッスン台本 フォーマット変更ツール")
    print("=" * 60)
    print("\n⚠️ 実行前に Word文書を閉じてください\n")

    format_specific_phrases(doc_path)
